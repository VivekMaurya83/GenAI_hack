import os
import sys
import json
import re
from typing import Optional, Tuple, List, Dict, Any

# Required libraries (ensure they are installed via requirements.txt)
import fitz  # PyMuPDF
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# Removed firebase_admin imports as they are used in db_core, not ai_core

# =========================
# Setup
# =========================
def setup_api():
    """Loads environment variables and configures the API key."""
    load_dotenv()
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key: raise ValueError("GOOGLE_API_KEY not found in your .env file.")
        genai.configure(api_key=api_key)
        return "gemini-1.5-flash-latest"
    except Exception as e:
        print(f"Error: API configuration failed. {e}")
        sys.exit(1)

MODEL_NAME = setup_api()

# =========================
# Helper Functions
# =========================
def _safe_json_loads(s: str, fallback=None):
    """Safely loads a JSON string, even if it's embedded in markdown."""
    if not s: return fallback
    s = s.strip()
    # Attempt to strip markdown code blocks
    if s.startswith("```json"): s = s[7:]
    if s.endswith("```"): s = s[:-3]
    s = s.strip() # Strip again after removing markdown

    try:
        return json.loads(s)
    except json.JSONDecodeError:
        # If direct parse fails, try to find a JSON-like substring
        m = re.search(r"\{.*\}", s, flags=re.DOTALL)
        if m:
            try: return json.loads(m.group(0))
            except json.JSONDecodeError: return fallback
    return fallback

def _norm(s: Optional[str]) -> str:
    """Returns a stripped string or an empty string if input is None."""
    return (s or "").strip()

def _smart_join(parts: List[Optional[str]]) -> str:
    """Joins a list of parts with a separator, ignoring empty or None parts."""
    return " | ".join([str(p) for p in parts if _norm(p)])

def _best_section_key(target_key: str, available_keys: List[str]) -> Optional[str]:
    """Finds the best matching key in a dictionary from a fuzzy user input."""
    if not target_key: return None
    t = target_key.strip().lower().replace(" ", "_")
    for k in available_keys:
        k_norm = k.lower().replace(" ", "_")
        if t == k_norm or t in k_norm or k_norm in t: return k
    return None

def parse_user_optimization_input(inp: str) -> Tuple[Optional[str], Optional[str]]:
    """Parses user input into a (section, instruction) tuple."""
    val = (inp or "").strip()
    if not val: return None, None
    if ":" in val:
        left, right = val.split(":", 1); return _norm(left), _norm(right)
    if len(val.split()) == 1:
        return val, None
    return None, val

def _stringify_list_content(content: Any) -> str: # Keep this here as save_resume_json_to_docx also uses it
    """Safely converts a list of strings or dicts into a single newline-separated string."""
    if not isinstance(content, list): return str(content or "")
    string_parts = []
    for item in content:
        if isinstance(item, str): string_parts.append(item)
        elif isinstance(item, dict):
            string_parts.append(", ".join([f"{k.replace('_', ' ').title()}: {v}" for k, v in item.items()]))
        else: string_parts.append(str(item))
    return "\n".join(string_parts)

def extract_text_auto(path: str) -> Optional[str]:
    if not os.path.exists(path):
        print(f"Error: File not found at path: {path}"); return None
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            with fitz.open(path) as doc: return "\n".join([page.get_text() for page in doc])
        elif ext == ".docx":
            doc = Document(path)
            chunks = [p.text for p in doc.paragraphs if _norm(p.text)]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells if _norm(cell.text)]
                    if cells: chunks.append(" | ".join(cells))
            return "\n".join(chunks)
        else:
            print(f"Error: Unsupported file type '{ext}'."); return None
    except Exception as e:
        print(f"Error reading file '{path}': {e}"); return None

def get_resume_structure(resume_text: str) -> Optional[Dict[str, Any]]:
    prompt = f"""
You are an expert HR Technology engineer specializing in resume data extraction. Your task is to convert the raw text of a resume into a structured, valid JSON object, capturing ALL information with high fidelity.
**Instructions:**
1.  **Use the Base Schema:** For common sections, use the following schema.
2.  **Capture Everything Else:** If you find other sections that do not fit the schema (e.g., "Achievements", "Leadership"), create a new top-level key for them (e.g., "achievements").
3.  **IGNORE THE SKILLS SECTION:** Do not parse the skills section in this step. It will be handled by a different process. Omit the 'skills' key from your output.
**Base Schema:**
{{
  "personal_info": {{ "name": "string", "email": "string", "phone": "string", "linkedin": "string", "github": "string" }},
  "summary": "string",
  "work_experience": [ {{ "role": "string", "company": "string", "duration": "string", "description": ["string", ...] }} ],
  "internships": [ {{ "role": "string", "company": "string", "duration": "string", "description": ["string", ...] }} ],
  "education": [ {{ "institution": "string", "degree": "string", "duration": "string", "description": ["string", ...] }} ],
  "projects": [ {{ "title": "string", "description": ["string", ...] }} ],
  "certifications": [ {{ "name": "string", "description": "string" }} ]
}}
**Critical Rules:**
- If a section from the base schema is NOT in the resume, YOU MUST OMIT ITS KEY from the final JSON. Do not create empty sections.
- Your final output must be a single, valid JSON object starting with `{{` and ending with `}}`. Do not include markdown.
--- RESUME TEXT ---
{resume_text}
--- END RESUME TEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        data = _safe_json_loads(response.text, fallback=None)
        if not data:
            print("\n--- ERROR: GEMINI API FAILED TO RETURN VALID JSON (STRUCTURE) ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-----------------------------------------------------\n")
            return None
        return data
    except Exception as e:
        print(f"Error during API call to Gemini: {e}"); return None

def categorize_skills_from_text(resume_text: str) -> Optional[Dict[str, List[str]]]:
    prompt = f"""
You are an expert technical recruiter and data analyst.
Your sole job is to scan the entire resume text provided and identify all skills, both technical and soft.
**Instructions:**
1.  Extract skills from *anywhere* in the text: summaries, project descriptions, a dedicated skills section, etc.
2.  Categorize the skills into the predefined keys in the JSON schema below.
3.  Place each skill only in the most appropriate category.
4.  If a category has no skills, you can omit the key from the output.
**JSON Output Schema:**
{{
    "Programming Languages": ["Python", "JavaScript", "Java", "C++", ...],
    "Frameworks and Libraries": ["TensorFlow", "PyTorch", "React", "Node.js", "Pandas", ...],
    "Databases": ["MySQL", "PostgreSQL", "MongoDB", ...],
    "Tools and Platforms": ["Git", "Docker", "AWS", "Jira", "Linux", ...],
    "Data Science": ["Machine Learning", "NLP", "Data Visualization", "Predictive Modeling", ...],
    "Soft Skills": ["Leadership", "Teamwork", "Communication", "Problem Solving", ...]
}}
**Critical Rules:**
- Your output must be ONLY the valid JSON object described above.
- Do not add any explanation or markdown.
--- RESUME TEXT ---
{resume_text}
--- END RESUME TEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        data = _safe_json_loads(response.text, fallback=None)
        if not data:
            print("\n--- ERROR: GEMINI FAILED TO INFER SKILLS ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-------------------------------------------------\n")
            return None
        return data
    except Exception as e:
        print(f"Error inferring skills with Gemini: {e}"); return None

def optimize_resume_json(resume_json: Dict[str, Any], user_input: str) -> Dict[str, Any]:
    section_req, instruction = parse_user_optimization_input(user_input)
    keys_present = list(resume_json.keys())
    model = genai.GenerativeModel(MODEL_NAME)
    base_prompt_context = f"""
CONTEXT: You are an elite career strategist and executive resume writer. Your task is to transform a resume from a passive list of duties into a compelling narrative of achievements that will impress top-tier recruiters.
**Your Transformation Checklist (Apply to every relevant bullet point):**
1.  **Lead with a Powerful Action Verb:** Replace weak verbs with strong, specific verbs (e.g., "Engineered," "Architected," "Spearheaded").
2.  **Quantify Metrics Relentlessly:** Add concrete numbers to show scale and achievement.
3.  **Showcase Impact and Scope:** If a number isn't available, describe the tangible impact or business outcome.
4.  **Integrate Technical Skills Naturally:** Weave technologies into the story of the achievement.
5.  **Ensure Brevity and Clarity:** Remove filler words. Each bullet point should be a single, powerful line.
**Critical Rules:**
- **Do not modify, add, or delete any titles, names, companies, institutions, or skill names.** This is a strict rule. Only rewrite descriptions.
- DO NOT invent facts or skills.
- DO NOT invent specific numbers.
- Preserve the original data structure.
- Do not modify personal information (name, email, phone).
- Your final output must be only the requested, valid JSON. Do not include markdown.
"""
    if section_req:
        mapped = _best_section_key(section_req, keys_present)
        if not mapped:
            print(f"⚠️ Section '{section_req}' not found."); return resume_json
        sec_data = resume_json.get(mapped)
        instr_text = instruction or "Apply your transformation checklist to make this section world-class."
        prompt = f"""
{base_prompt_context}
TASK: Apply your full transformation checklist to optimize ONLY the following JSON section, named "{mapped}".
--- INPUT JSON SECTION ---
{json.dumps(sec_data, indent=2)}
--- END INPUT JSON SECTION ---
"""
    else:
        prompt = f"""
{base_prompt_context}
TASK: Apply your full transformation checklist to optimize all sections of the following resume JSON.
--- FULL INPUT JSON ---
{json.dumps(resume_json, indent=2)}
--- END INPUT JSON ---
"""
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        optimized_data = _safe_json_loads(response.text, fallback=None)
        
        if not optimized_data:
            print("\n--- ERROR: GEMINI API FAILED TO RETURN VALID JSON (OPTIMIZE) ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-----------------------------------------------------\n")
            return resume_json
            
        if section_req and optimized_data:
            resume_json[mapped] = optimized_data
        elif optimized_data:
            for key, value in optimized_data.items():
                if key in resume_json and isinstance(resume_json[key], dict) and isinstance(value, dict):
                    resume_json[key].update(value)
                elif key in resume_json and isinstance(resume_json[key], list) and isinstance(value, list):
                    resume_json[key] = value
                else:
                    resume_json[key] = value
    except Exception as e:
        print(f"Error during optimization: {e}")
    return resume_json

def optimize_for_linkedin(resume_json: Dict[str, Any], user_input: str) -> Optional[Dict[str, Any]]:
    context_text = []
    if 'summary' in resume_json: context_text.append(f"Summary:\n{resume_json['summary']}")
    
    all_experiences = resume_json.get('work_experience', []) + resume_json.get('internships', [])
    if all_experiences:
        context_text.append("\nProfessional Experience & Internships:")
        for job in all_experiences:
            description_str = ' '.join(job.get('description', []) if isinstance(job.get('description'), list) else [str(job.get('description', ''))])
            context_text.append(f"- {job.get('role')} at {job.get('company')}: {description_str}")
    
    if 'projects' in resume_json:
        context_text.append("\nProjects:")
        for project in resume_json['projects']:
            description_str = ' '.join(project.get('description', []) if isinstance(project.get('description'), list) else [str(project.get('description', ''))])
            context_text.append(f"- {project.get('title')}: {description_str}")
    
    if 'skills' in resume_json and isinstance(resume_json['skills'], dict):
        skills_summary = ", ".join([f"{cat}: {', '.join(skills)}" for cat, skills in resume_json['skills'].items()])
        context_text.append(f"\nSkills: {skills_summary}")
    
    for key, value in resume_json.items():
        if key not in ['personal_info', 'summary', 'work_experience', 'internships', 'projects', 'skills', 'education', 'certifications', 'resume_metadata']:
            if isinstance(value, str):
                context_text.append(f"\n{key.replace('_', ' ').title()}:\n{value}")
            elif isinstance(value, list):
                context_text.append(f"\n{key.replace('_', ' ').title()}:\n" + "\n".join([str(item) for item in value]))

    resume_context = "\n".join(context_text)
    section_req, instruction = parse_user_optimization_input(user_input)
    base_prompt_context = f"""
You are an expert LinkedIn profile strategist and personal branding coach.
Your task is to generate compelling, optimized text for a user's LinkedIn profile based on the provided resume content.
**Instructions:**
1.  **Headlines:** Create 2-3 powerful, keyword-rich headline options.
2.  **About (Summary):** Write a compelling, first-person "About" section.
3.  **Experiences:** For EACH job/internship in the context, rewrite the bullet points to be concise and results-oriented.
4.  **Projects:** For EACH project in the context, rewrite its description to be engaging for a LinkedIn audience.
**JSON Output Schema:**
{{
    "headlines": ["string option 1", ...],
    "about_section": "string",
    "optimized_experiences": [ {{ "title": "Role at Company", "description": "string" }} ],
    "optimized_projects": [ {{ "title": "Project Title", "description": "string" }} ]
}}
**Critical Rules:**
- Generate content ONLY from the provided resume context.
- Keep the tone professional but approachable.
- Your final output must be ONLY the valid JSON object that matches the requested task.
"""
    if section_req:
        instr_text = instruction or f"Make the {section_req} section more compelling and professional."
        prompt = f"""
{base_prompt_context}
TASK: Based on the resume context, optimize ONLY the '{section_req}' portion of a LinkedIn profile.
--- RESUME CONTEXT ---
{resume_context}
--- END RESUME CONTEXT ---
"""
    else:
        instr_text = instruction or "Optimize the entire LinkedIn profile, processing every experience and project."
        prompt = f"""
{base_prompt_context}
TASK: Based on the resume context, perform a full optimization of a LinkedIn profile.
--- RESUME CONTEXT ---
{resume_context}
--- END RESUME CONTEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        data = _safe_json_loads(response.text, fallback=None)
        if not data:
            print("\n--- ERROR: GEMINI FAILED TO INFER LINKEDIN CONTENT ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-------------------------------------------------\n")
            return None
        return data
    except Exception as e:
        print(f"Error generating LinkedIn content with Gemini: {e}"); return None

def save_resume_json_to_docx(resume_json: Dict[str, Any]) -> Document:
    doc = Document()
    def add_heading(text: Optional[str], level: int = 1):
        t = _norm(text);
        if t: h = doc.add_heading(t, level=level); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def add_para(text: Optional[str], bold: bool = False, style: Optional[str] = None):
        t = _norm(text)
        if t:
            p = doc.add_paragraph(style=style)
            run = p.add_run(t)
            run.bold = bold
            run.font.size = Pt(11)
            if style == "List Bullet": p.paragraph_format.left_indent = Pt(36)
            
    print_order = ['personal_info', 'summary', 'skills', 'work_experience', 'internships', 'projects', 'education', 'certifications']
    
    name_for_title = resume_json.get('personal_info', {}).get('name', '')
    if name_for_title:
        doc.add_heading(name_for_title, level=0)

    contact_info_parts = []
    p_info = resume_json.get('personal_info', {})
    if p_info.get('email'): contact_info_parts.append(p_info['email'])
    if p_info.get('phone'): contact_info_parts.append(p_info['phone'])
    if p_info.get('linkedin'): contact_info_parts.append(p_info['linkedin'])
    if p_info.get('github'): contact_info_parts.append(p_info['github'])
    if contact_info_parts:
        add_para(_smart_join(contact_info_parts))
    
    for section in print_order:
        if section in resume_json:
            content = resume_json[section]
            if section == 'personal_info':
                continue
            
            add_heading(section.replace("_", " ").title(), level=2)
            
            if section == 'summary' and isinstance(content, str):
                add_para(content)
            elif section == 'skills' and isinstance(content, dict):
                for category, skill_list in content.items():
                    if isinstance(skill_list, list) and skill_list:
                        p = doc.add_paragraph();
                        run = p.add_run(category.replace("_", " ").title() + ': '); run.bold = True
                        p.add_run(", ".join(skill_list)); run.font.size = Pt(11)
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, str):
                        add_para(item, style="List Bullet")
                    elif isinstance(item, dict):
                        header_parts = [item.get("title"), item.get("name"), item.get("role"), item.get("degree"), item.get("institution")]
                        header = _smart_join(header_parts)
                        if header: add_para(header, bold=True)
                        
                        sub_header_parts = [item.get("company"), item.get("duration")]
                        sub_header = _smart_join(sub_header_parts)
                        if sub_header: add_para(sub_header)
                        
                        desc = item.get("description", [])
                        if isinstance(desc, list):
                            for bullet in desc:
                                if _norm(bullet): add_para(str(bullet), style="List Bullet")
                        elif _norm(desc):
                            add_para(str(desc), style="List Bullet")
            elif isinstance(content, str):
                add_para(content)
            
    for section, content in resume_json.items():
        if section not in print_order and section not in ['resume_metadata']:
            add_heading(section.replace("_", " ").title(), level=2)
            if isinstance(content, list):
                for item in content: add_para(str(item), style="List Bullet")
            else: add_para(str(content))
            
    print("\n✅ DOCX document generated in memory.")
    return doc

def generate_career_roadmap(user_profile: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    prompt = f"""
    Act as a world-class AI Career Strategist and Technical Project Manager. Your task is to generate a deeply personalized, multi-faceted career action plan.

    **STEP 1: ANALYZE THE USER'S PROFILE**
    - **User's Current State (from resume or manual input):** ```{user_profile.get('current_skills_input')}```
    - **Stated Current Proficiency:** {user_profile.get('current_level')}
    - **User's Stated Goal (Job Description or Desired Skills):** ```{user_profile.get('goal_input')}```
    - **Desired Goal Proficiency:** {user_profile.get('goal_level')}
    - **Time Commitment:** Plan for a duration of **{user_profile.get('duration')}**, assuming **{user_profile.get('study_hours')}** study hours per month.

    **STEP 2: GENERATE THE ACTION PLAN AS A SINGLE, VALID JSON OBJECT**
    The JSON output must be perfectly structured with the following keys. Do not include any explanatory text outside of the JSON object.

    1.  "domain": A single string representing the most relevant domain inferred from the goal input (e.g., "Data Science", "Cybersecurity"). This is a new, crucial key.
    2.  "extracted_skills_and_projects": A JSON object with "skills" (array of strings) and "projects" (array of strings).
    3.  "job_match_score": A JSON object with "score" (number) and "summary" (string).
    4.  "skills_to_learn_summary": An array of strings.
    5.  "timeline_chart_data": A JSON object with "labels" (array of strings) and "durations" (array of numbers in weeks) ans also the total weeks counts must be equal to users specified duration.
    6.  "detailed_roadmap": An array of "phase" objects, each with "phase_title", "phase_duration", and "topics" (array of strings).
    7.  "suggested_projects": An array of 2 "project" objects, each with "project_title", "project_level", "skills_mapped", "what_you_will_learn", and a multi-step "implementation_plan".
    8.  "suggested_courses": THIS IS A CRITICAL SECTION.
        - You MUST generate an array of 2-3 "course" objects.
        - Each object MUST contain the following FOUR keys: "course_name", "platform", "url", and "mapping".
        - The "platform" MUST be a string like "Coursera", "edX", "Pluralsight", etc.
        - The "url" MUST be a direct, fully-qualified, and workable hyperlink.
        - The "mapping" MUST be a concise sentence explaining how the course helps the roadmap.
        - **Follow this example format precisely:**
          `{{ "course_name": "Google Data Analytics Certificate", "platform": "Coursera", "url": "https://www.coursera.org/professional-certificates/google-data-analytics", "mapping": "This certificate covers the foundational skills in Phase 1 and 2." }}`
    """
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        cleaned_response_text = response.text.replace('```json', '').replace('```', '').strip()
        return json.loads(cleaned_response_text)
    except Exception as e:
        print(f"An error occurred during AI roadmap generation: {e}"); return None

def get_tutor_explanation(topic: str) -> Optional[Dict[str, Any]]:
    prompt = f"""
    Act as a friendly and encouraging expert tutor. A user is currently working through a personalized learning plan and is stuck on the following topic: **"{topic}"**

    Your task is to provide a clear, helpful explanation in a structured JSON format. The JSON object must have the following keys:

    1.  **"analogy"**: A simple, real-world analogy to help the user understand the core concept intuitively.
    2.  **"technical_definition"**: A concise, technically accurate definition. If the topic involves code, provide a short, well-commented code snippet in the appropriate language (e.g., Python, JavaScript).
    3.  **"prerequisites"**: An array of 1-3 prerequisite concepts the user might need to review. This helps them identify foundational knowledge gaps.

    Generate the JSON object and nothing else.
    """
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        cleaned_response_text = response.text.replace('```json', '').replace('```', '').strip()
        return json.loads(cleaned_response_text)
    except Exception as e:
        print(f"An error occurred in AI Tutor: {e}"); return None

def get_chatbot_response(query: str, history: List[Dict[str, str]], career_plan: Dict[str, Any]) -> Optional[Dict[str, str]]:
    formatted_history = [{'role': msg['role'], 'parts': [{'text': msg['content']}]} for msg in history]
    
    career_plan_details = "No career plan provided."
    if career_plan:
        plan = career_plan
        roadmap_topics = [topic for phase in plan.get('detailed_roadmap', []) for topic in phase.get('topics', [])]
        
        career_plan_details = (
            f"**Domain:** {plan.get('domain', 'N/A')}\n\n"
            f"**Job Match Score:** {plan.get('job_match_score', {}).get('score', 'N/A')}% with a summary of: \"{plan.get('job_match_score', {}).get('summary', 'N/A')}\"\n\n"
            f"**Priority Skills to Learn:** {', '.join(plan.get('skills_to_learn_summary', []))}\n\n"
            f"**Roadmap Topics:** {', '.join(roadmap_topics)}\n\n"
            f"**Projects:**\n"
            f"1. {plan.get('suggested_projects', [{}])[0].get('project_title', 'N/A')}\n"
            f"2. {plan.get('suggested_projects', [{}, {}])[1].get('project_title', 'N/A')}\n\n"
            f"**Suggested Courses:**\n"
            f"1. {plan.get('suggested_courses', [{}])[0].get('course_name', 'N/A')}\n"
            f"2. {plan.get('suggested_courses', [{}, {}])[1].get('course_name', 'N/A')}\n"
        )
    
    system_instruction_prompt = (
        f"You are an AI career strategist and tutor. Your purpose is to provide concise, point-to-point, and beginner-friendly guidance to the user, strictly based on the career plan provided below.\n\n"
        f"**Career Plan Details:**\n{career_plan_details}\n\n"
        f"**Your Instructions:**\n"
        f"1. Keep responses brief, beginner-friendly, and to the point.\n"
        f"2. You can answer questions related to the provided career plan, including the **job match score, priority skills, timeline, detailed roadmap, projects, and courses**.\n"
        f"3. If the user asks a question that is **outside the scope** of the career plan's domain or is not directly related to the provided plan data, you must respond with a polite refusal. For example, 'That question seems to be outside the scope of your current career plan. Is there anything I can help you with related to your career plan?'\n\n"
        f"Let's begin."
    )
    
    formatted_history.insert(0, {'role': 'model', 'parts': [{'text': "Okay, I'm ready to help based on the user's career plan."}]})
    formatted_history.insert(0, {'role': 'user', 'parts': [{'text': system_instruction_prompt}]})

    model = genai.GenerativeModel(MODEL_NAME)
    try:
        chat_session = model.start_chat(history=formatted_history)
        response = chat_session.send_message(
            query,
            generation_config=genai.types.GenerationConfig(temperature=0.4),
        )
        return {"response": response.text}
    except Exception as e:
        print(f"An error occurred in the chat endpoint: {e}"); return None