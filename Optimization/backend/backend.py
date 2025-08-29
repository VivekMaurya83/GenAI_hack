import os
import sys
import json
import re
from typing import Optional, Tuple

# Required libraries (ensure they are installed via requirements.txt)
import fitz  # PyMuPDF
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import firebase_admin
from firebase_admin import credentials, firestore

# =========================
# Setup
# =========================
def setup_api():
    """Loads environment variables and configures the API key."""
    load_dotenv()
    try:
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key: raise ValueError("GOOGLE_API_KEY not found in your .env file.")
        genai.configure(api_key=api_key)
        return "gemini-1.5-flash-latest"
    except Exception as e:
        print(f"Error: API configuration failed. {e}")
        sys.exit(1)

MODEL_NAME = setup_api()

# =========================
# Helper Functions
# =========================
def _safe_json_loads(s: str, fallback=None):
    """Safely loads a JSON string, even if it's embedded in markdown."""
    if not s: return fallback
    s = s.strip()
    if s.startswith("```json"): s = s[7:]
    if s.endswith("```"): s = s[:-3]
    try: return json.loads(s)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", s, flags=re.DOTALL)
        if m:
            try: return json.loads(m.group(0))
            except json.JSONDecodeError: return fallback
    return fallback

def _norm(s: Optional[str]) -> str:
    """Returns a stripped string or an empty string if input is None."""
    return (s or "").strip()

def _smart_join(parts):
    """Joins a list of parts with a separator, ignoring empty or None parts."""
    return " | ".join([str(p) for p in parts if _norm(p)])

def _best_section_key(target_key: str, available_keys):
    """Finds the best matching key in a dictionary from a fuzzy user input."""
    if not target_key: return None
    t = target_key.strip().lower().replace(" ", "_")
    for k in available_keys:
        k_norm = k.lower().replace(" ", "_")
        if t == k_norm or t in k_norm or k_norm in t: return k
    return None

def parse_user_optimization_input(inp: str) -> Tuple[Optional[str], Optional[str]]:
    """Parses user input into a (section, instruction) tuple."""
    val = (inp or "").strip()
    if not val: return None, None
    if ":" in val:
        left, right = val.split(":", 1); return _norm(left), _norm(right)
    if len(val.split()) == 1:
        return val, None
    return None, val

def _stringify_list_content(content) -> str:
    """Safely converts a list of strings or dicts into a single newline-separated string."""
    if not isinstance(content, list): return str(content or "")
    string_parts = []
    for item in content:
        if isinstance(item, str): string_parts.append(item)
        elif isinstance(item, dict):
            string_parts.append(", ".join([f"{k.replace('_', ' ').title()}: {v}" for k, v in item.items()]))
        else: string_parts.append(str(item))
    return "\n".join(string_parts)

# =========================
# Database Manager Class
# =========================
class DatabaseManager:
    """Handles all interactions with the Firebase Firestore database."""
    def __init__(self):
        try:
            if not firebase_admin._apps:
                cred = credentials.Certificate("firebase-credentials.json")
                firebase_admin.initialize_app(cred)
            self.db = firestore.client()
            print("✅ Firebase connection successful.")
        except Exception as e:
            print(f"❌ Firebase Error: {e}")
            print("Ensure 'firebase-credentials.json' is in the 'backend' folder.")
            sys.exit(1)

    def _map_ai_section_to_standard_key(self, ai_key: str) -> Optional[str]:
        normalized_key = ai_key.lower().replace(" ", "_").replace("-", "_")
        mapping = {
            'work_experience': ['work_experience', 'professional_experience', 'experience', 'work_history'],
            'education': ['education', 'academic_background'], 'projects': ['projects', 'personal_projects'],
            'internships': ['internships', 'internship_experience'],
            'certifications': ['certifications', 'licenses_&_certifications']
        }
        for standard_key, variations in mapping.items():
            if normalized_key in variations: return standard_key
        return None

    def save_resume_relational(self, file_name: str, parsed_data: dict) -> str:
        p_info = parsed_data.get('personal_info', {})
        resume_doc_ref = self.db.collection('resumes').document()
        resume_doc_ref.set({
            'file_name': file_name, 'name': p_info.get('name'), 'email': p_info.get('email'),
            'phone': p_info.get('phone'), 'linkedin': p_info.get('linkedin'),
            'github': p_info.get('github'), 'summary': parsed_data.get('summary'),
            'optimized_summary': None
        })
        resume_id = resume_doc_ref.id
        print(f" -> Saved base resume to Firestore with ID: {resume_id}")

        known_sections_db_map = {
            'work_experience': 'work_experiences', 'education': 'education',
            'projects': 'projects', 'internships': 'internships',
            'certifications': 'certifications'
        }
        for ai_section_key, section_content in parsed_data.items():
            if ai_section_key in ['personal_info', 'summary', 'skills']: continue
            standard_key = self._map_ai_section_to_standard_key(ai_section_key)
            if standard_key:
                table_name = known_sections_db_map[standard_key]
                if isinstance(section_content, list):
                    for item in section_content:
                        if isinstance(item, dict):
                            item['description'] = _stringify_list_content(item.get('description'))
                            item['optimized_description'] = None
                            resume_doc_ref.collection(table_name).add(item)
            else:
                description = _stringify_list_content(section_content)
                resume_doc_ref.collection('additional_sections').add({
                    'section_name': ai_section_key,
                    'description': description,
                    'optimized_description': None
                })
        if 'skills' in parsed_data and isinstance(parsed_data['skills'], dict):
            for category, skill_list in parsed_data['skills'].items():
                if isinstance(skill_list, list):
                    for skill_name in skill_list:
                        resume_doc_ref.collection('skills').add({'category': category, 'skill_name': skill_name})
        return resume_id

    def fetch_resume_relational(self, resume_id: str, get_optimized: bool = False) -> Optional[dict]:
        resume_doc_ref = self.db.collection('resumes').document(resume_id)
        resume_doc = resume_doc_ref.get()
        if not resume_doc.exists: return None
        main_info = resume_doc.to_dict()
        resume_data = {}
        resume_data['personal_info'] = {k: main_info.get(k) for k in ['name', 'email', 'phone', 'linkedin', 'github']}
        resume_data['summary'] = (main_info.get('optimized_summary') if get_optimized and main_info.get('optimized_summary') else main_info.get('summary'))
        known_collections_map = {
            'work_experience': ('work_experiences', ['role', 'company', 'duration']),
            'education': ('education', ['institution', 'degree', 'duration']),
            'projects': ('projects', ['title']),
            'internships': ('internships', ['role', 'company', 'duration']),
            'certifications': ('certifications', ['name'])
        }
        for key, (collection_name, columns) in known_collections_map.items():
            docs = resume_doc_ref.collection(collection_name).stream()
            data_list = []
            for doc in docs:
                item_data = doc.to_dict()
                desc_to_use = (item_data.get('optimized_description') if get_optimized and item_data.get('optimized_description') else item_data.get('description'))
                item = {col: item_data.get(col) for col in columns}
                if desc_to_use: item['description'] = desc_to_use.split('\n')
                data_list.append(item)
            if data_list: resume_data[key] = data_list
        docs = resume_doc_ref.collection('skills').stream()
        skills_dict = {}
        for doc in docs:
            item = doc.to_dict()
            category = item.get('category'); skill_name = item.get('skill_name')
            if category and skill_name:
                if category not in skills_dict: skills_dict[category] = []
                skills_dict[category].append(skill_name)
        if skills_dict: resume_data['skills'] = skills_dict
        docs = resume_doc_ref.collection('additional_sections').stream()
        for doc in docs:
            item = doc.to_dict()
            desc_to_use = (item.get('optimized_description') if get_optimized and item.get('optimized_description') else item.get('description'))
            section_name = item.get('section_name')
            if section_name and desc_to_use: resume_data[section_name] = desc_to_use.split('\n')
        return {k: v for k, v in resume_data.items() if v}

    def update_optimized_resume_relational(self, resume_id: str, optimized_data: dict):
        resume_doc_ref = self.db.collection('resumes').document(resume_id)
        if 'summary' in optimized_data: resume_doc_ref.update({'optimized_summary': optimized_data['summary']})
        def update_many(collection_name: str, items: list, match_keys: list):
            for item in items:
                desc = _stringify_list_content(item.get('description', []))
                query = resume_doc_ref.collection(collection_name)
                for key in match_keys:
                    if item.get(key): query = query.where(key, '==', item.get(key))
                docs = query.limit(1).stream()
                for doc in docs: doc.reference.update({'optimized_description': desc})
        if 'work_experience' in optimized_data: update_many('work_experiences', optimized_data['work_experience'], ['role', 'company'])
        if 'education' in optimized_data: update_many('education', optimized_data['education'], ['institution', 'degree'])
        if 'projects' in optimized_data: update_many('projects', optimized_data['projects'], ['title'])
        if 'internships' in optimized_data: update_many('internships', optimized_data['internships'], ['role', 'company'])
        if 'certifications' in optimized_data: update_many('certifications', optimized_data['certifications'], ['name'])
        for key, content in optimized_data.items():
            if self._map_ai_section_to_standard_key(key) is None and key not in ['personal_info', 'summary', 'skills']:
                desc = _stringify_list_content(content)
                docs = resume_doc_ref.collection('additional_sections').where('section_name', '==', key).limit(1).stream()
                for doc in docs: doc.reference.update({'optimized_description': desc})
        print(f" -> Optimized data for resume ID {resume_id} has been fully updated in Firestore.")

    def close_connection(self): pass

def extract_text_auto(path: str) -> Optional[str]:
    if not os.path.exists(path):
        print(f"Error: File not found at path: {path}"); return None
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            with fitz.open(path) as doc: return "\n".join([page.get_text() for page in doc])
        elif ext == ".docx":
            doc = Document(path)
            chunks = [p.text for p in doc.paragraphs if _norm(p.text)]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells if _norm(cell.text)]
                    if cells: chunks.append(" | ".join(cells))
            return "\n".join(chunks)
        else:
            print(f"Error: Unsupported file type '{ext}'."); return None
    except Exception as e:
        print(f"Error reading file '{path}': {e}"); return None

def get_resume_structure(resume_text: str) -> Optional[dict]:
    prompt = f"""
You are an expert HR Technology engineer specializing in resume data extraction. Your task is to convert the raw text of a resume into a structured, valid JSON object, capturing ALL information with high fidelity.
**Instructions:**
1.  **Use the Base Schema:** For common sections, use the following schema.
2.  **Capture Everything Else:** If you find other sections that do not fit the schema (e.g., "Achievements", "Leadership"), create a new top-level key for them (e.g., "achievements").
3.  **IGNORE THE SKILLS SECTION:** Do not parse the skills section in this step. It will be handled by a different process. Omit the 'skills' key from your output.
**Base Schema:**
{{
  "personal_info": {{ "name": "string", "email": "string", "phone": "string", "linkedin": "string", "github": "string" }},
  "summary": "string",
  "work_experience": [ {{ "role": "string", "company": "string", "duration": "string", "description": ["string", ...] }} ],
  "internships": [ {{ "role": "string", "company": "string", "duration": "string", "description": ["string", ...] }} ],
  "education": [ {{ "institution": "string", "degree": "string", "duration": "string", "description": ["string", ...] }} ],
  "projects": [ {{ "title": "string", "description": ["string", ...] }} ],
  "certifications": [ {{ "name": "string", "description": "string" }} ]
}}
**Critical Rules:**
- If a section from the base schema is NOT in the resume, YOU MUST OMIT ITS KEY from the final JSON. Do not create empty sections.
- Your final output must be a single, valid JSON object starting with `{{` and ending with `}}`. Do not include markdown.
--- RESUME TEXT ---
{resume_text}
--- END RESUME TEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        data = _safe_json_loads(response.text, fallback=None)
        if not data:
            print("\n--- ERROR: GEMINI API FAILED TO RETURN VALID JSON (STRUCTURE) ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-----------------------------------------------------\n")
            return None
        return data
    except Exception as e:
        print(f"Error during API call to Gemini: {e}"); return None

def categorize_skills_from_text(resume_text: str) -> Optional[dict]:
    prompt = f"""
You are an expert technical recruiter and data analyst.
Your sole job is to scan the entire resume text provided and identify all skills, both technical and soft.
**Instructions:**
1.  Extract skills from *anywhere* in the text: summaries, project descriptions, a dedicated skills section, etc.
2.  Categorize the skills into the predefined keys in the JSON schema below.
3.  Place each skill only in the most appropriate category.
4.  If a category has no skills, you can omit the key from the output.
**JSON Output Schema:**
{{
    "Programming Languages": ["Python", "JavaScript", "Java", "C++", ...],
    "Frameworks and Libraries": ["TensorFlow", "PyTorch", "React", "Node.js", "Pandas", ...],
    "Databases": ["MySQL", "PostgreSQL", "MongoDB", ...],
    "Tools and Platforms": ["Git", "Docker", "AWS", "Jira", "Linux", ...],
    "Data Science": ["Machine Learning", "NLP", "Data Visualization", "Predictive Modeling", ...],
    "Soft Skills": ["Leadership", "Teamwork", "Communication", "Problem Solving", ...]
}}
**Critical Rules:**
- Your output must be ONLY the valid JSON object described above.
- Do not add any explanation or markdown.
--- RESUME TEXT ---
{resume_text}
--- END RESUME TEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        data = _safe_json_loads(response.text, fallback=None)
        if not data:
            print("\n--- ERROR: GEMINI FAILED TO INFER SKILLS ---")
            print("API Response Text:", response.text)
            try: print("API Prompt Feedback:", response.prompt_feedback)
            except ValueError: pass
            print("-------------------------------------------------\n")
            return None
        return data
    except Exception as e:
        print(f"Error inferring skills with Gemini: {e}"); return None

def optimize_resume_json(resume_json: dict, user_input: str) -> dict:
    section_req, instruction = parse_user_optimization_input(user_input)
    keys_present = set(resume_json.keys())
    model = genai.GenerativeModel(MODEL_NAME)
    base_prompt_context = f"""
CONTEXT: You are an elite career strategist and executive resume writer. Your task is to transform a resume from a passive list of duties into a compelling narrative of achievements that will impress top-tier recruiters.
**Your Transformation Checklist (Apply to every relevant bullet point):**
1.  **Lead with a Powerful Action Verb:** Replace weak verbs with strong, specific verbs (e.g., "Engineered," "Architected," "Spearheaded").
2.  **Quantify Metrics Relentlessly:** Add concrete numbers to show scale and achievement.
3.  **Showcase Impact and Scope:** If a number isn't available, describe the tangible impact or business outcome.
4.  **Integrate Technical Skills Naturally:** Weave technologies into the story of the achievement.
5.  **Ensure Brevity and Clarity:** Remove filler words. Each bullet point should be a single, powerful line.
**Critical Rules:**
- **Do not modify, add, or delete any titles, names, companies, institutions, or skill names.** This is a strict rule. Only rewrite descriptions.
- DO NOT invent facts or skills.
- DO NOT invent specific numbers.
- Preserve the original data structure.
- Do not modify personal information (name, email, phone).
- Your final output must be only the requested, valid JSON. Do not include markdown.
"""
    if section_req:
        mapped = _best_section_key(section_req, keys_present)
        if not mapped:
            print(f"⚠️ Section '{section_req}' not found."); return resume_json
        sec_data = resume_json.get(mapped)
        instr_text = instruction or "Apply your transformation checklist to make this section world-class."
        prompt = f"""
{base_prompt_context}
TASK: Apply your full transformation checklist to optimize ONLY the following JSON section, named "{mapped}".
--- INPUT JSON SECTION ---
{json.dumps(sec_data, indent=2)}
--- END INPUT JSON SECTION ---
"""
    else:
        prompt = f"""
{base_prompt_context}
TASK: Apply your full transformation checklist to optimize all sections of the following resume JSON.
--- FULL INPUT JSON ---
{json.dumps(resume_json, indent=2)}
--- END FULL INPUT JSON ---
"""
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        optimized_data = _safe_json_loads(response.text, fallback=None)
        if section_req and optimized_data: resume_json[mapped] = optimized_data
        elif optimized_data: resume_json = optimized_data
    except Exception as e: print(f"Error during optimization: {e}")
    return resume_json

def optimize_for_linkedin(resume_json: dict, user_input: str) -> Optional[dict]:
    context_text = []
    if 'summary' in resume_json: context_text.append(f"Summary:\n{resume_json['summary']}")
    all_experiences = resume_json.get('work_experience', []) + resume_json.get('internships', [])
    if all_experiences:
        context_text.append("\nProfessional Experience & Internships:")
        for job in all_experiences: context_text.append(f"- {job.get('role')} at {job.get('company')}: {' '.join(job.get('description', []))}")
    if 'projects' in resume_json:
        context_text.append("\nProjects:")
        for project in resume_json['projects']: context_text.append(f"- {project.get('title')}: {' '.join(project.get('description', []))}")
    if 'skills' in resume_json:
        skills_summary = ", ".join([f"{cat}: {', '.join(skills)}" for cat, skills in resume_json['skills'].items()])
        context_text.append(f"\nSkills: {skills_summary}")
    resume_context = "\n".join(context_text)
    section_req, instruction = parse_user_optimization_input(user_input)
    base_prompt_context = f"""
You are an expert LinkedIn profile strategist and personal branding coach.
Your task is to generate compelling, optimized text for a user's LinkedIn profile based on the provided resume content.
**Instructions:**
1.  **Headlines:** Create 2-3 powerful, keyword-rich headline options.
2.  **About (Summary):** Write a compelling, first-person "About" section.
3.  **Experiences:** For EACH job/internship in the context, rewrite the bullet points to be concise and results-oriented.
4.  **Projects:** For EACH project in the context, rewrite its description to be engaging for a LinkedIn audience.
**JSON Output Schema:**
{{
    "headlines": ["string option 1", ...],
    "about_section": "string",
    "optimized_experiences": [ {{ "title": "Role at Company", "description": "string" }} ],
    "optimized_projects": [ {{ "title": "Project Title", "description": "string" }} ]
}}
**Critical Rules:**
- Generate content ONLY from the provided resume context.
- Keep the tone professional but approachable.
- Your final output must be ONLY the valid JSON object that matches the requested task.
"""
    if section_req:
        instr_text = instruction or f"Make the {section_req} section more compelling and professional."
        prompt = f"""
{base_prompt_context}
TASK: Based on the resume context, optimize ONLY the '{section_req}' portion of a LinkedIn profile.
--- RESUME CONTEXT ---
{resume_context}
--- END RESUME CONTEXT ---
"""
    else:
        instr_text = instruction or "Optimize the entire LinkedIn profile, processing every experience and project."
        prompt = f"""
{base_prompt_context}
TASK: Based on the resume context, perform a full optimization of a LinkedIn profile.
--- RESUME CONTEXT ---
{resume_context}
--- END RESUME CONTEXT ---
"""
    model = genai.GenerativeModel(MODEL_NAME)
    try:
        safety_settings = {'HARM_CATEGORY_HARASSMENT': 'BLOCK_NONE','HARM_CATEGORY_HATE_SPEECH': 'BLOCK_NONE','HARM_CATEGORY_SEXUALLY_EXPLICIT': 'BLOCK_NONE','HARM_CATEGORY_DANGEROUS_CONTENT': 'BLOCK_NONE'}
        response = model.generate_content(prompt, safety_settings=safety_settings)
        return _safe_json_loads(response.text, fallback=None)
    except Exception as e:
        print(f"Error generating LinkedIn content with Gemini: {e}"); return None

def save_resume_json_to_docx(resume_json: dict) -> Document:
    doc = Document()
    def add_heading(text, level=1):
        t = _norm(text);
        if t: h = doc.add_heading(t, level=level); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def add_para(text, bold=False):
        t = _norm(text)
        if t: p = doc.add_paragraph(); run = p.add_run(t); run.bold = bold; run.font.size = Pt(11)
    print_order = ['personal_info', 'summary', 'skills', 'work_experience', 'internships', 'projects', 'education', 'certifications']
    for section in print_order:
        if section in resume_json:
            content = resume_json[section]
            if section == 'personal_info':
                name = content.get('name', '');
                if name: add_heading(name, level=0)
                contact_info = [content.get('email'), content.get('phone'), content.get('linkedin'), content.get('github')]
                add_para(_smart_join(contact_info))
                continue
            add_heading(section.replace("_", " ").title(), level=2)
            if section == 'skills':
                for category, skill_list in content.items():
                    if isinstance(skill_list, list):
                        p = doc.add_paragraph(); p.add_run(category.replace("_", " ").title() + ': ').bold = True; p.add_run(", ".join(skill_list))
                continue
            if isinstance(content, str): add_para(content)
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, str): doc.add_paragraph(item, style="List Bullet")
                    elif isinstance(item, dict):
                        header_parts = [item.get("title"), item.get("name"), item.get("role"), item.get("degree"), item.get("institution")]
                        header = _smart_join(header_parts)
                        if header: add_para(header, bold=True)
                        sub_header_parts = [item.get("company"), item.get("duration")]
                        sub_header = _smart_join(sub_header_parts)
                        if sub_header: add_para(sub_header)
                        desc = item.get("description", [])
                        if isinstance(desc, list):
                            for bullet in desc:
                                if _norm(bullet): doc.add_paragraph(str(bullet), style="List Bullet")
                        elif _norm(desc): doc.add_paragraph(str(desc), style="List Bullet")
    for section, content in resume_json.items():
        if section not in print_order:
            add_heading(section.replace("_", " ").title(), level=2)
            if isinstance(content, list):
                for item in content: doc.add_paragraph(str(item), style="List Bullet")
            else: add_para(str(content))
    print("\n✅ DOCX document generated in memory.")
    return doc
